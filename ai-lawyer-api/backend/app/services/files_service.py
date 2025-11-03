# -*- coding: utf-8 -*-

from typing import Optional, Tuple, List, Dict, Any
from pathlib import Path
from io import BytesIO
from html.parser import HTMLParser
import os
import logging
import re
import tempfile
import subprocess
import shutil

from ..db.db import get_pg_pool
from ..db.repositories import files_repo


logger = logging.getLogger(__name__)


# ------------------ 自定义异常 ------------------

class FileExtractError(Exception):
    pass


class UnsupportedFileType(FileExtractError):
    pass


class DependencyMissing(FileExtractError):
    pass


# ------------------ 文本抽取 ------------------

def _extract_docx_text(data: bytes) -> str:
    try:
        import docx  # python-docx
    except Exception as e:
        raise DependencyMissing(f"python-docx 加载失败: {e}")
    try:
        document = docx.Document(BytesIO(data))
        paras = [p.text for p in document.paragraphs if p.text]
        return "\n".join(paras).strip()
    except Exception as e:
        raise FileExtractError(f"DOCX 文本抽取失败: {e}")


def _extract_pdf_text(data: bytes) -> str:
    reader = None
    try:
        import pypdf  # type: ignore
        reader = pypdf.PdfReader(BytesIO(data))
    except Exception:
        try:
            import PyPDF2  # type: ignore
            reader = PyPDF2.PdfReader(BytesIO(data))
        except Exception as e:
            raise DependencyMissing(f"pypdf/PyPDF2 加载失败: {e}")
    try:
        texts: List[str] = []
        for page in reader.pages:
            t = page.extract_text() or ""
            if t:
                texts.append(t)
        return "\n\n".join(texts).strip()
    except Exception as e:
        raise FileExtractError(f"PDF 文本抽取失败: {e}")


def _is_text_like(filename: str, content_type: Optional[str]) -> bool:
    ct = (content_type or "").lower()
    if ct.startswith("text/"):
        return True
    if ct in {"application/json", "application/xml"}:
        return True
    suffix = Path(filename or "").suffix.lower()
    return suffix in {".txt", ".md", ".json", ".csv", ".log"}


def extract_file_text(data: bytes, filename: str, content_type: Optional[str]) -> str:
    """从上传数据中提取文本（不落盘）。
    仅支持 docx/pdf/文本类文件，不支持图片、加密 PDF 等。
    失败时抛出 FileExtractError 及其子类，由调用方映射 HTTP 状态码。
    """
    if not data:
        raise FileExtractError("空文件")
    suffix = Path(filename or "").suffix.lower()
    ct = (content_type or "").lower()

    if suffix == ".docx" or "word" in ct:
        return _extract_docx_text(data)
    if suffix == ".pdf" or ct == "application/pdf":
        return _extract_pdf_text(data)
    if _is_text_like(filename, content_type):
        try:
            return data.decode("utf-8", errors="ignore").strip()
        except Exception:
            raise FileExtractError("文本解码失败")

    raise UnsupportedFileType("仅支持 docx、pdf、文本类文件")


# ------------------ 文件表 CRUD ------------------

async def list_files_service(
    *,
    name: Optional[str] = None,
    user_id: Optional[int] = None,
    page: int = 1,
    page_size: int = 20,
) -> Tuple[List[Dict[str, Any]], int]:
    """分页查询文件列表，可按名称和用户过滤"""
    pool = await get_pg_pool()
    skip = (page - 1) * page_size
    items = await files_repo.get_all(
        pool,
        skip=skip,
        limit=page_size,
        name=name,
        user_id=user_id,
    )
    total = await files_repo.count(pool, name=name, user_id=user_id)
    return items, total


async def get_file_by_id(file_id: int) -> Optional[Dict[str, Any]]:
    pool = await get_pg_pool()
    return await files_repo.get_by_id(pool, file_id)


async def create_file(data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    pool = await get_pg_pool()
    new_id = await files_repo.create(pool, data)
    return await files_repo.get_by_id(pool, new_id)


async def update_file(file_id: int, data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    pool = await get_pg_pool()
    ok = await files_repo.update(pool, file_id, data)
    if not ok:
        return None
    return await files_repo.get_by_id(pool, file_id)


async def delete_file(file_id: int) -> bool:
    pool = await get_pg_pool()
    return await files_repo.delete(pool, file_id)


# ------------------ 存储路径与安全 ------------------

def _files_root() -> Path:
    return Path(os.environ.get("FILES_ROOT", "/app/files")).resolve()


def _safe_join_file_path(doc_url: str) -> Path:
    """规范映射：数据库存储以 "files/" 开头；物理位置为 FILES_ROOT/去前缀路径。"""
    files_root = _files_root()
    rel_raw = (doc_url or "").strip()
    rel = rel_raw.replace("\\", "/").lstrip("/")
    if not rel.lower().startswith("files/"):
        raise FileExtractError("非法文件路径")
    rel_no_prefix = rel[6:]
    abs_path = (files_root / rel_no_prefix).resolve()
    if not str(abs_path).startswith(str(files_root)):
        raise FileExtractError("非法文件路径")
    if not abs_path.exists():
        raise FileExtractError("文件不存在")
    return abs_path


# ------------------ HTML 渲染辅助 ------------------

def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _docx_to_html(data: bytes) -> str:
    """优先使用 mammoth 生成更保真的 HTML，失败再回退到 python-docx 的简易方案。"""
    try:
        import mammoth  # type: ignore
        try:
            result = mammoth.convert_to_html(BytesIO(data))
            html = result.value or ""
            return f"<div>{html}</div>"
        except Exception:
            pass
    except Exception:
        pass

    try:
        import docx
    except Exception as e:
        raise DependencyMissing(f"python-docx 加载失败: {e}")
    try:
        document = docx.Document(BytesIO(data))
        html_parts: List[str] = ["<div>"]
        for p in document.paragraphs:
            style_name = (getattr(p.style, 'name', '') or '').lower()
            if style_name.startswith('heading'):
                level = ''.join(ch for ch in style_name if ch.isdigit()) or '1'
                tag = f"h{level}"
            else:
                tag = "p"
            runs_html: List[str] = []
            for r in p.runs:
                t = _escape_html(r.text or "")
                if not t:
                    continue
                open_tags = ""
                close_tags = ""
                if r.bold:
                    open_tags += "<strong>"
                    close_tags = "</strong>" + close_tags
                if r.italic:
                    open_tags += "<em>"
                    close_tags = "</em>" + close_tags
                if r.underline:
                    open_tags += "<u>"
                    close_tags = "</u>" + close_tags
                runs_html.append(f"{open_tags}{t}{close_tags}")
            inner = ''.join(runs_html) or "<br/>"
            html_parts.append(f"<{tag}>{inner}</{tag}>")
        html_parts.append("</div>")
        return "".join(html_parts)
    except Exception as e:
        raise FileExtractError(f"DOCX 转 HTML 失败: {e}")


def _text_to_html(text: str) -> str:
    parts: List[str] = ["<div>"]
    for para in (text or "").splitlines():
        parts.append(f"<p>{_escape_html(para)}</p>")
    parts.append("</div>")
    return "".join(parts)

 
def _pdf_to_html(data: bytes) -> str:
    text = _extract_pdf_text(data)
    return _text_to_html(text)


# ------------------ LibreOffice 转换 ------------------

def _run_soffice_convert(input_path: Path, out_dir: Path, to: str) -> bool:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False
    profile_dir = out_dir / "lo_profile"
    cmd = [
        soffice,
        "--headless",
        f"-env:UserInstallation=file:///{profile_dir.as_posix()}",
        "--convert-to",
        to,
        "--outdir",
        str(out_dir),
        str(input_path),
    ]
    try:
        timeout_s = 180 if to.startswith("docx") else 120
        proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout_s)
        if proc.returncode != 0:
            logger.warning("LibreOffice convert failed rc=%s to=%s stderr=%s", proc.returncode, to, proc.stderr.decode(errors="ignore"))
            return False
        return True
    except Exception:
        logger.exception("LibreOffice convert exception to=%s", to)
        return False


def _try_doc_to_html_via_libreoffice(data: bytes) -> str | None:
    with tempfile.TemporaryDirectory() as td:
        in_path = Path(td) / "input.doc"
        out_dir = Path(td)
        in_path.write_bytes(data)
        if not _run_soffice_convert(in_path, out_dir, "html:HTML"):
            if not _run_soffice_convert(in_path, out_dir, "html"):
                return None
        out_html = out_dir / "input.html"
        if out_html.exists():
            return out_html.read_text(encoding="utf-8", errors="ignore")
        return None


def _try_docx_to_html_via_libreoffice(data: bytes) -> str | None:
    with tempfile.TemporaryDirectory() as td:
        in_path = Path(td) / "input.docx"
        out_dir = Path(td)
        in_path.write_bytes(data)
        if not _run_soffice_convert(in_path, out_dir, "html:HTML"):
            if not _run_soffice_convert(in_path, out_dir, "html"):
                return None
        out_html = out_dir / "input.html"
        if out_html.exists():
            return out_html.read_text(encoding="utf-8", errors="ignore")
        return None


def _doc_to_html_via_libreoffice_mammoth(data: bytes) -> str | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    try:
        import mammoth  # type: ignore
    except Exception:
        return None
    with tempfile.TemporaryDirectory() as td:
        in_path = Path(td) / "input.doc"
        out_dir = Path(td)
        in_path.write_bytes(data)
        if not _run_soffice_convert(in_path, out_dir, "docx"):
            return None
        out_docx = out_dir / "input.docx"
        if not out_docx.exists():
            return None
        with out_docx.open("rb") as f:
            result = mammoth.convert_to_html(f)
            html = result.value or ""
            return f"<div>{html}</div>"


# ------------------ 样式与清理 ------------------

def _inject_default_css(html: str) -> str:
    style = (
        "<style>"
        "body,div{font-family: system-ui,Segoe UI,Roboto,Helvetica,Arial,'Noto Sans SC',sans-serif;}"
        "p{margin:0.6em 0;line-height:1.7;}"
        "p:has(>br:only-child){display:none;}"
        "h1{font-size:1.8em;margin:0.8em 0 0.4em;}"
        "h2{font-size:1.6em;margin:0.8em 0 0.4em;}"
        "h3{font-size:1.4em;margin:0.8em 0 0.4em;}"
        "ul,ol{margin:0.6em 1.4em;}"
        "table{border-collapse:collapse;margin:0.6em 0;width:100%;}"
        "th,td{border:1px solid #ddd;padding:6px;vertical-align:top;}"
        "</style>"
    )
    if "<style" in html:
        return html
    if html.lstrip().startswith("<div"):
        return html.replace("<div>", f"<div>{style}", 1)
    return style + html




# ------------------ 对外：HTML 预览 ------------------

async def get_file_as_html(file_id: int, mode: str = "semantic") -> str:
    obj = await get_file_by_id(file_id)
    if not obj:
        raise FileExtractError("文件不存在")
    abs_path = _safe_join_file_path(obj.get("doc_url") or "")
    if not abs_path.exists():
        raise FileExtractError("文件不存在")
    suffix = abs_path.suffix.lower()
    data = abs_path.read_bytes()

    if suffix == ".docx":
        # 强制走 mammoth 转换路径
        try:
            import mammoth  # type: ignore
        except Exception as e:
            raise DependencyMissing(f"mammoth 加载失败: {e}")
        with open(abs_path, "rb") as f:
            result = mammoth.convert_to_html(f)
        html = result.value or ""
        return html

    if suffix == ".pdf":
        html = _pdf_to_html(data)
        return _inject_default_css(html) if mode == "semantic" else html

    if suffix == ".doc":
        # 先转 docx 再用 mammoth 转 html
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            raise DependencyMissing("缺少 LibreOffice (soffice) 用于将 .doc 转为 .docx")
        try:
            import mammoth  # type: ignore
        except Exception as e:
            raise DependencyMissing(f"mammoth 加载失败: {e}")
        with tempfile.TemporaryDirectory() as td:
            in_path = abs_path
            out_dir = Path(td)
            # 转换为 docx
            if not _run_soffice_convert(in_path, out_dir, "docx"):
                raise FileExtractError(".doc 转 .docx 失败")
            docx_path = out_dir / "input.docx"
            if not docx_path.exists():
                # 某些版本会保留原文件名
                alts = list(out_dir.glob("*.docx"))
                if not alts:
                    raise FileExtractError("未找到转换后的 .docx 文件")
                docx_path = alts[0]
            with open(docx_path, "rb") as f:
                result = mammoth.convert_to_html(f)
            html = result.value or ""
            return _inject_default_css(f"<div>{html}</div>")

    # 其余文本类：utf-8 解码并包装为 HTML
    try:
        text = data.decode("utf-8", errors="ignore")
        html = _text_to_html(text)
        return _inject_default_css(html) if mode == "semantic" else html
    except Exception:
        raise UnsupportedFileType("暂不支持该文件格式的 HTML 预览")


# ------------------ HTML -> DOCX 写回 ------------------

class _HTMLToDocxParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.blocks: List[Dict[str, Any]] = []  # 每个 block: {tag, runs: [{text,bold,italic,underline}]}
        self.current_runs: List[Dict[str, Any]] = []
        self.style_stack: List[Dict[str, bool]] = []
        self.current_block_tag: str = "p"

    def _push_block(self, tag: str):
        if self.current_runs:
            self.blocks.append({"tag": self.current_block_tag, "runs": self.current_runs})
            self.current_runs = []
        self.current_block_tag = tag

    def handle_starttag(self, tag, attrs):
        t = tag.lower()
        if t in ("p", "div", "h1", "h2", "h3", "h4", "h5", "h6"):
            self._push_block(t if t != "div" else "p")
        elif t == "br":
            self.current_runs.append({"text": "\n", "bold": False, "italic": False, "underline": False})
        style = {"bold": False, "italic": False, "underline": False}
        if t in ("strong", "b"):
            style["bold"] = True
        if t in ("em", "i"):
            style["italic"] = True
        if t == "u":
            style["underline"] = True
        if self.style_stack:
            parent = self.style_stack[-1].copy()
            for k, v in style.items():
                parent[k] = parent.get(k, False) or v
            self.style_stack.append(parent)
        else:
            self.style_stack.append(style)

    def handle_endtag(self, tag):
        t = tag.lower()
        if t in ("p", "div", "h1", "h2", "h3", "h4", "h5", "h6"):
            if self.current_runs:
                self.blocks.append({"tag": self.current_block_tag, "runs": self.current_runs})
                self.current_runs = []
            self.current_block_tag = "p"
        if self.style_stack:
            self.style_stack.pop()

    def handle_data(self, data):
        text = data or ""
        if not text:
            return
        style = self.style_stack[-1] if self.style_stack else {"bold": False, "italic": False, "underline": False}
        self.current_runs.append({
            "text": text,
            "bold": bool(style.get("bold")),
            "italic": bool(style.get("italic")),
            "underline": bool(style.get("underline")),
        })


def _html_to_docx_bytes(html: str) -> bytes:
    try:
        import docx
    except Exception as e:
        raise DependencyMissing(f"python-docx 加载失败: {e}")
    parser = _HTMLToDocxParser()
    parser.feed(html or "")
    document = docx.Document()
    for block in (parser.blocks or [{"tag": "p", "runs": [{"text": html or ""}]}]):
        tag = block.get("tag") or "p"
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            p = document.add_paragraph()
            try:
                p.style = f"Heading {tag[1:]}"
            except Exception:
                pass
        else:
            p = document.add_paragraph()
        for r in block.get("runs", []):
            run = p.add_run(r.get("text", ""))
            run.bold = r.get("bold", False)
            run.italic = r.get("italic", False)
            run.underline = r.get("underline", False)
    bio = BytesIO()
    document.save(bio)
    return bio.getvalue()


async def update_file_with_html(file_id: int, html: str) -> Optional[Dict[str, Any]]:
    obj = await get_file_by_id(file_id)
    if not obj:
        return None
    files_root = _files_root()
    rel = (obj.get("doc_url") or "").replace("\\", "/").lstrip("/")
    if not rel.lower().startswith("files/"):
        raise FileExtractError("非法文件路径")
    rel_no_prefix = rel[6:]
    old_abs = (files_root / rel_no_prefix).resolve()
    if not str(old_abs).startswith(str(files_root)):
        raise FileExtractError("非法文件路径")
    if old_abs.suffix.lower() == ".docx":
        target_abs = old_abs
        new_rel = str(Path("files") / rel_no_prefix).replace("\\", "/")
        new_name = obj.get("name") or old_abs.name
    else:
        target_abs = old_abs.with_suffix('.docx')
        new_rel = str(Path("files") / Path(rel_no_prefix).with_suffix('.docx')).replace("\\", "/")
        stem = (obj.get("name") or old_abs.name)
        new_name = str(Path(stem).with_suffix('.docx'))

    data = _html_to_docx_bytes(html or "")
    target_abs.parent.mkdir(parents=True, exist_ok=True)
    target_abs.write_bytes(data)

    if target_abs != old_abs and old_abs.exists():
        try:
            old_abs.unlink()
        except Exception:
            pass

    updated = await update_file(file_id, {
        "doc_url": new_rel,
        "name": new_name,
    })
    return updated
